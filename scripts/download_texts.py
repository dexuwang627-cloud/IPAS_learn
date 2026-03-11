#!/usr/bin/env python3
"""Download text-based files from IPAS Google Drive folder and extract text."""
import subprocess
import json
import os
import shutil

GWS = os.path.expanduser("~/.cargo/bin/gws")
OUTPUT_DIR = "data/texts"
TEMP_DIR = "data/_temp"

# Files to download: (file_id, filename, chapter)
FILES = [
    # V1 AI彙整資料 (AI-generated summaries, ~2.5MB PDFs with text layer)
    ("1cvLfmo_vznQg5_wkWFLffl-_6jlqO2Dj", "L211_組織節能減碳策略_V1整理.pdf", "L211 組織節能減碳策略"),
    ("1UeawFLvT-hysOgDnKOmstgPdFXA0aEGU", "L212_節能技術應用與能源管理_V1整理.pdf", "L212 節能技術應用與能源管理"),
    ("1TfbJKIJE2Rr63DXZSOHc2ncVLg5CSJJ7", "L213_再生能源與綠電導入_V1整理.pdf", "L213 再生能源與綠電導入"),
    ("10aA2K0xX4guxn3kKbvOeC-V6D8CC32Pi", "學習指南_V1整理.pdf", "學習指南"),

    # 考試彙整 docx files (study group exam summaries)
    ("1aKFV93q9MM6SGzvdgD9tRT5EOCoefhgy", "L21202_節能技術投資_黃文弘.docx", "L21202 節能技術投資效益"),
    ("1acz-7W_AZ7MEISTdqxB3uiD49UYSqhFW", "L21201_公用設施節能_鄭惠如.docx", "L21201 公用設施節能技術"),
    ("1XWfJh3cmi-bcjdyxUkqo6CwXO528PRH6", "L21204_ISO50001_大瑋.docx", "L21204 ISO 50001能源管理"),
    ("1CEFLMDML7w3aZ9HkIpWLQ3r8Zd2GyrJ0", "L21203_ESCO_劉怡麟.docx", "L21203 ESCO應用實務"),
    ("1anBx-G5WBKfXJfMzAEGmsW5bGrs-DP98", "L21205_節能前瞻技術_Emily.docx", "L21205 節能相關前瞻技術"),
    ("1Qf9hq9zXeAexORsSD5VGsoiTrEhw4HLi", "第一組讀書會.docx", "讀書會筆記"),

    # 重要計算題型
    ("1foSgyGeofXycj3enD5RJLECdTXdhIHY1", "碳管理制度實務_計算題型1.docx", "碳管理制度實務"),
    ("1s2xT0ENZPneRQSNFJeYDpne5TcJsOTcA", "碳管理制度實務_計算題型2.docx", "碳管理制度實務"),
    ("14fu464QSvUFmejxI2x2alQPPQ1aImMN1", "碳管理制度實務_計算題型.pdf", "碳管理制度實務"),

    # 說明會簡報 + 重點報告
    ("1XTezuiGnj8Pu9T3ppNzGSDgppSZeFN3i", "iPAS說明會簡報.pdf", "IPAS說明會"),
    ("1Lj2ZjjcU_2eqFlZkkipzvIk31MUCh-xv", "讀書會重點報告_第二組.pdf", "讀書會筆記"),

    # 前瞻能源
    ("1TWKkqdVZL_N7MeHAZb269wjMmNBs4eOE", "臺灣2050淨零轉型_前瞻能源.pdf", "前瞻能源"),
]


def download_file(file_id: str, dest_path: str, base_dir: str) -> bool:
    """Download a file from Google Drive using gws CLI."""
    temp_dir = os.path.join(base_dir, TEMP_DIR)
    os.makedirs(temp_dir, exist_ok=True)

    try:
        result = subprocess.run(
            [GWS, "drive", "files", "get", "--params",
             json.dumps({"fileId": file_id, "alt": "media"})],
            capture_output=True, text=True, timeout=120,
            cwd=temp_dir
        )
        # gws saves to download.pdf or similar in cwd
        for f in os.listdir(temp_dir):
            if f.startswith("download"):
                shutil.move(os.path.join(temp_dir, f), dest_path)
                return True
        return False
    except Exception as e:
        print(f"  Error: {e}")
        return False


def extract_text_from_docx(path: str) -> str:
    """Extract text from a .docx file."""
    from docx import Document
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    return "\n".join(paragraphs)


def extract_text_from_pdf(path: str) -> str:
    """Extract text from a PDF file."""
    import pdfplumber
    try:
        with pdfplumber.open(path) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n\n".join(pages)
    except Exception:
        return ""


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(TEMP_DIR, exist_ok=True)
    base_dir = os.path.abspath(".")

    results = {}
    for file_id, filename, chapter in FILES:
        print(f"Downloading: {filename}...", end=" ", flush=True)

        ext = os.path.splitext(filename)[1]
        temp_path = os.path.join(base_dir, TEMP_DIR, filename)

        if download_file(file_id, os.path.join(base_dir, TEMP_DIR, filename), base_dir):
            # Extract text
            if ext == ".docx":
                text = extract_text_from_docx(temp_path)
            elif ext == ".pdf":
                text = extract_text_from_pdf(temp_path)
            else:
                text = ""

            if text.strip():
                txt_name = filename.rsplit(".", 1)[0] + ".txt"
                out_path = os.path.join(OUTPUT_DIR, txt_name)
                with open(out_path, "w", encoding="utf-8") as f:
                    f.write(f"# Chapter: {chapter}\n\n{text}")
                results[filename] = {"chars": len(text), "chapter": chapter}
                print(f"OK ({len(text)} chars)")
            else:
                print("NO TEXT (skipped)")

            # Clean up temp file
            if os.path.exists(temp_path):
                os.remove(temp_path)
        else:
            print("DOWNLOAD FAILED")

    # Save summary
    with open(os.path.join(OUTPUT_DIR, "_summary.json"), "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    print(f"\nDone! Extracted {len(results)} files to {OUTPUT_DIR}")

    # Clean up
    if os.path.exists(TEMP_DIR):
        shutil.rmtree(TEMP_DIR)


if __name__ == "__main__":
    main()
